## @file WorshipList.py
#  @brief Generates a worship chart from specifies songs and keys.
#  @author Samuel Crawford
#  @date 1/23/2019

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT

class FileError(Exception):
    pass

class ParamError(Exception):
    pass

def main():
    doc = Document()
    
    # Defines margins

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    doc = writeSong(doc, "ThePassion", "D")
    doc = writeSong(doc, "Anointing", "B")
    doc = writeSong(doc, "OPraiseTheName", "B")
    doc = writeSong(doc, "DeathWasArrested", "B")
    
    #doc = writeSong(doc, "GloriousDay", "D")
    #doc = writeSong(doc, "WhoYouSayIAm", "F#")
    #doc = writeSong(doc, "LetThereBeLight", "C")
    #doc = writeSong(doc, "SpiritOfTheLivingGod", "B")

    doc.save('output.docx')

def writeSong(doc, fileName, oldKey):
    infile = open("songs/" + fileName + ".txt", "r")
    lines = infile.readlines()
    infile.close()

    key = oldKey[0].upper()
    if len(oldKey) > 1:
        key += oldKey[1]

    # Defines default style

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(28)

    # Writes title

    p = doc.add_paragraph()
    title = p.add_run(lines[0].strip() + " ")
    title.font.size = Pt(36)
    title.underline = True

    titleKey = p.add_run("(" + key + ")")
    titleKey.font.size = Pt(20)
    titleKey.underline = True

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    # Writes chords

    for i in lines[1:]:
        line = i.split()

        p = doc.add_paragraph()

        # Defines tab stops
        
        tab_stops = p.paragraph_format.tab_stops
        tab_stop = tab_stops.add_tab_stop(Inches(1.58), WD_TAB_ALIGNMENT.LEFT)

        if line[0][-1] == ":":
            p.add_run(line[0] + "\t")
            chordStart = 1
        else:
            p.add_run(line[0] + " " + line[1] + "\t")
            chordStart = 2

        for chord in line[chordStart:]:
            if chord == "|":
                p.add_run("|  ")
            elif chord == "new":
                p.add_run("\n\t")
            elif chord == "double":
                p.add_run("x 2  ")
            elif chord == "triple":
                p.add_run("x 3  ")
            elif "/" in chord:
                newChord = chord[:-1]
                p.add_run(getChord(key, newChord, fileName) + "/")
            elif "(" in chord:
                newChord = chord[1:]
                p.add_run("(" + getChord(key, newChord, fileName) + "  ")
            elif ")" in chord:
                newChord = chord[:-1]
                p.add_run(getChord(key, newChord, fileName) + ")  ")
            else:
                p.add_run(getChord(key, chord, fileName) + "  ")

        p.paragraph_format.line_spacing = Pt(36)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    return doc

def getNotes(key, notes):
    noteList = []
    note = notes.index(key)
    for i in list(range(7)):
        noteList.append(notes[note])
        if i == 2:
            note += 1
        else:
            note += 2
    return noteList

def getChord(key, num, fileName):
    validKeys  = ['F','F#','Gb','G','G#','Ab','A','A#','Bb','B','C','C#','Db','D','D#','Eb','E']
    notesSharp = ['F','F#','G','G#','A','A#','B','C','C#','D','D#','E','F','F#','G','G#','A','A#','B','C','C#','D','D#','E']
    notesFlat  = ['F','Gb','G','Ab','A','Bb','B','C','Db','D','Eb','E','F','Gb','G','Ab','A','Bb','B','C','Db','D','Eb','E']

    if key not in validKeys:
        raise ParamError("The key \"" + key + "\" isn't recognized.")

    if len(key) > 1:
        if key[1] == "#":
            keyList = getNotes(key, notesSharp)
        elif key[1] == "b":
            keyList = getNotes(key, notesFlat)
    elif key in ["C", "F"]:
        keyList = getNotes(key, notesFlat)
    else:
        keyList = getNotes(key, notesSharp)        

    minor = False
    if num.islower():
        minor = True

    num = num.lower()

    # TODO: make more efficient
    if num == "i":
        chord = keyList[0]
    elif num == "ii":
        chord = keyList[1]
    elif num == "iii":
        chord = keyList[2]
    elif num == "iv":
        chord = keyList[3]
    elif num == "v":
        chord = keyList[4]
    elif num == "vi":
        chord = keyList[5]
    elif num == "vii":
        chord = keyList[6]
    else:
        raise FileError("The chord \"" + num + "\" in the " + fileName + " file isn't recognized.")
    if minor:
        chord += "m"

    return chord

main()