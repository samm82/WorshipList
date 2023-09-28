## @file   Document.py
#  @brief  Contains functions for adding text to document.
#  @author Samuel Crawford
#  @date   9/28/2023

import win32com.client

from enum import Enum, auto
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT

from Helpers import getChord, getNotes


## @brief          Outputs a .pdf from a .docx file.
#  @param[in] docx The filename of the .docx file.
#  @param[in] pdf  The filename of the .pdf file.
#  @return         True if the conversion was successful and False otherwise.
def pdfWrite(docx, pdf):
    word = win32com.client.DispatchEx('Word.Application')
    success = False
    try:
        doc = word.Documents.Open(str(docx))
        doc.SaveAs(str(pdf), FileFormat=17)
        doc.Close()
        success = True
    except AttributeError:
        pass
    finally:
        word.Quit()
        return success


## @brief  Sets up an empty document.
#  @return The document.
def docSetup():
    doc = Document()

    # Defines margins

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Defines default style

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(28)

    return doc


## @brief               Writes a song to doc.
#  @param[in] doc       The document being generated.
#  @param[in] lineCount A counter of how many lines have been printed.
#  @param[in] fileName  The name of the song file.
#  @param[in] key       The key of the song.
#  @return              The document (doc) and line counter (lineCount).
def writeSong(doc, lineCount, fileName, key):
    with Path(f"src/songs/{fileName}.txt").open() as fp:
        lines = fp.readlines()

    # Adds page break if song will get cut off
    newLineLength = len(lines)
    for line in lines:
        if "new" in line.split():
            newLineLength += 1

    if lineCount + newLineLength > 15:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        lineCount = newLineLength
    else:
        lineCount += newLineLength

    # Writes title
    doc = writeTitle(doc, lines[0], key)

    # Gets list of notes from getNotes(key)
    noteList = getNotes(key)

    # Writes the lines with chords
    for i in range(1, len(lines)):
        line = lines[i].split()
        end = i == len(lines) - 1

        doc = writeLine(doc, line, end, noteList, fileName)

    return doc, lineCount


## @brief            Writes a song title to the document.
#  @param[in] doc    The document to write to.
#  @param[in] title  The title to write to the document.
#  @param[in] key    The key of the song.
#  @return           The document.
def writeTitle(doc, title, key):
    p = doc.add_paragraph()
    title = p.add_run(title.strip() + " ")
    title.font.size = Pt(36)
    title.underline = True

    titleKey = p.add_run(f"({key})")
    titleKey.font.size = Pt(20)
    titleKey.underline = True

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    return doc


## @brief          Writes a section name to the document.
#  @param[in] p    The paragraph to write to.
#  @param[in] line The line to get the section name from.
#  @param[in] sep  The separator to use after the section name.
#  @param[in] ind  The index to be read next.
#  @return         The paragraph and the next index.
def writeSection(p, line, sep, ind):
    items = [line[ind]]
    i = 0

    while line[ind + i][-1] != ":":
        i += 1
        items.append(line[ind + i])

    p.add_run(" ".join(items) + sep)
    return p, ind + len(items)


## @brief           Writes a line to the document.
#  @param[in] doc   The document to write to.
#  @param[in] line  The line to write to the document.
#  @param[in] end   True if line is the last line in the song and False otherwise.
#  @param[in] notes The list of valid notes.
#  @param[in] file  The file with the line to write.
#  @return          The document.
def writeLine(doc, line, end, notes, file):
    p = doc.add_paragraph()

    # Defines tab stops
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(1.58), WD_TAB_ALIGNMENT.LEFT)

    # Defines relative font size based on song file
    class Size(Enum):
        NORMAL = auto()
        SMALL = auto()
        SMALL_LAST = auto()

    size = Size.NORMAL

    # Index variable for chord
    p, i = writeSection(p, line, "\t", 0)

    # Adds all chords
    lineLength = len(line)
    while i != lineLength:
        chord = line[i]
        iNext = True
        if chord == "|":
            run = p.add_run("|")
        elif chord == "new":
            run = p.add_run("\n")
        elif chord == "same":
            run = p.add_run("|  ")
            run, i = writeSection(p, line, "  ", i + 1)
            iNext = False
        elif chord[0] == "x" and len(chord) > 1 and chord[1:].isdecimal():
            run = p.add_run(f"x{chord[1:].lstrip('0')}")
        elif chord[0] == "(" and chord[-1] == ")":
            run = p.add_run("(" + getChord(notes, chord[1:-1], file) + ")")
            size = Size.SMALL_LAST
        elif chord[0] == "(":
            run = p.add_run("(" + getChord(notes, chord[1:], file))
            size = Size.SMALL
        elif chord[-1] == ")":
            run = p.add_run(getChord(notes, chord[:-1], file) + ")")
            size = Size.SMALL_LAST
        else:
            run = p.add_run(getChord(notes, chord, file))

        # Set font size for small text
        if size in {Size.SMALL, Size.SMALL_LAST}:
            run.font.size = Pt(22)

        # Increments i if necessary
        if iNext:
            i += 1

        # Adds space after chord
        if i != lineLength:
            if chord == "new":
                run = p.add_run("\t")
            else:
                run = p.add_run("  ")

            # Set font size for small text
            if size == Size.SMALL:
                run.font.size = Pt(22)
            elif size == Size.SMALL_LAST:
                run.font.size = Pt(22)
                size = Size.NORMAL

    # Sets paragraph spacing
    if end:
        p.paragraph_format.space_after = Pt(10)
    else:
        p.paragraph_format.space_after = Pt(0)

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(36)

    return doc
