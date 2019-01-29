## @file   WorshipList.py
#  @brief  Generates a worship chart from specified songs and keys.
#  @author Samuel Crawford
#  @date   1/28/2019

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT

from sys import argv

from Document import docSetup, writeLine
from GUI import songGUI
from MusicData import getNotes

## @brief The main function of the program that calls other programs.
def main():
    doc = docSetup()
    lineCount = 0

    # Writes each song

    if "gui" in argv:
        songs, keys, file = songGUI()
        for i in range(4):
            doc, lineCount = writeSong(doc, lineCount, songs[i].replace(" ", ""), keys[i])
        doc.save(file + ".docx")
    else:
        doc, lineCount = writeSong(doc, lineCount, "ThePassion", "D")
        doc, lineCount = writeSong(doc, lineCount, "Anointing", "B")
        doc, lineCount = writeSong(doc, lineCount, "OPraiseTheName", "B")
        doc, lineCount = writeSong(doc, lineCount, "DeathWasArrested", "B")
        doc.save('output.docx')

## @brief               Writes a song to doc.
#  @param[in] doc       The document being generated.
#  @param[in] lineCount A counter of how many lines have been printed.
#  @param[in] fileName  The name of the song file.
#  @param[in] oldKey    The original key of the song (manipulated in function).
#  @return              The document (doc) and line counter (lineCount).
def writeSong(doc, lineCount, fileName, oldKey):
    infile = open("src/songs/" + fileName + ".txt", "r")
    lines = infile.readlines()
    infile.close()

    # Converts key to uppercase

    key = oldKey[0].upper()
    if len(oldKey) > 1:
        key += oldKey[1]

    # Adds page break if song will get cut off

    newLineLength = len(lines)
    for line in lines:
        if "new" in line.split():
            newLineLength += 1

    lineCount += newLineLength

    if lineCount > 16:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
    
    # Writes title

    p = doc.add_paragraph()
    title = p.add_run(lines[0].strip() + " ")
    title.font.size = Pt(36)
    title.underline = True

    titleKey = p.add_run("(" + key + ")")
    titleKey.font.size = Pt(20)
    titleKey.underline = True

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

    # Gets list of notes from getNotes(key)

    noteList = getNotes(key)

    # Writes the lines with chords

    for i in range(1, len(lines)):
        line = lines[i].split()
        end = i == len(lines) - 1

        doc, newLines = writeLine(doc, line, end, noteList, fileName)
        lineCount += newLines

    return doc, lineCount

main()